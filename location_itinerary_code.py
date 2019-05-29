import os
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from datetime import datetime as dt, timedelta

# Increase max_columns output to show all columns
pd.options.display.max_columns = None
# Increase max_rows output to show all rows
pd.options.display.max_rows = None
# Be able to read entire value in each column (no longer truncating values)
pd.set_option('display.max_colwidth', -1)
# See all columns in one row of printout in console
pd.set_option('expand_frame_repr', False)

##########################################
### Create weekly location itineraries ###
##########################################
def location_itineraries(name_of_user, work_location):
    # Set working directory to location itineraries
    os.chdir('/Users/dustinwicker/PyCharmProjects/location_itinerary/')

    # Create dict assigning number to day of week
    days_week_dict = dict(zip([0, 1, 2, 3, 4, 5, 6],
                              ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']))
    # Grab first 5 days of week
    days_week = [value for value in days_week_dict.values()][0:5]

    # Creaing on Monday
    if dt.today().date().weekday() == 0:
        print("Today is Monday. Creating report...")
        # Create month, day for current week (M/DD)
        month_day = [str((dt.today().date() + timedelta(days=i)).month) + '/' +\
                     str((dt.today().date() + timedelta(days=i)).day) for i in range (7, 12, 1)]
    # Creaing on Tuesday
    if dt.today().date().weekday() == 1:
        print("Today is Tuesday. Creating report...")
        # Create month, day for current week (M/DD)
        month_day = [str((dt.today().date() + timedelta(days=i)).month) + '/' + \
                         str((dt.today().date() + timedelta(days=i)).day) for i in range(6, 11, 1)]
    # Creaing on Wednesday
    if dt.today().date().weekday() == 2:
        print("Today is Wednesday. Creating report...")
        # Create month, day for current week (M/DD)
        month_day = [str((dt.today().date() + timedelta(days=i)).month) + '/' + \
                         str((dt.today().date() + timedelta(days=i)).day) for i in range(5, 10, 1)]
    # Creaing on Thursday
    if dt.today().date().weekday() == 3:
        print("Today is Thursday. Creating report...")
        # Create month, day for current week (M/DD)
        month_day = [str((dt.today().date() + timedelta(days=i)).month) + '/' + \
                         str((dt.today().date() + timedelta(days=i)).day) for i in range(4, 9, 1)]
    # Creaing on Friday
    if dt.today().date().weekday() == 4:
        print("Today is Friday. Creating report...")
        # Create month, day for current week (M/DD)
        month_day = [str((dt.today().date() + timedelta(days=i)).month) + '/' + \
                         str((dt.today().date() + timedelta(days=i)).day) for i in range(3, 8, 1)]

    # Combine day of week with M/DD
    month_day_days_week = [days_week[i] + ',  ' + month_day[i] for i in range(0, 5)]
    # Set work location list
    work_location = work_location
    # Create DataFrame from lists, transpose, and rename columns as necessary
    location_itinerary_df = pd.DataFrame(data=[month_day_days_week, work_location]).transpose().\
        rename(columns= {0: 'Date', 1: 'Working Location'})

    # Create Document object
    document = Document()
    # Grab 'Normal' style object
    style = document.styles['Normal']
    # Set style.font to font
    font = style.font
    # Set font name
    font.name = 'Calibri'
    # Set font point size
    font.size = Pt(11)
    # Create paragraph object
    p = document.add_paragraph()
    # Append run to paragraph and underline
    p.add_run('Itinerary - ' + name_of_user).underline = True
    # Create align run
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Create paragraph object
    p = document.add_paragraph()
    # Append run to paragraph
    p.add_run('Week of ' + month_day[0] + ' - ' + month_day[-1])
    # Center align paragraph
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Add table to document
    t = document.add_table(rows=location_itinerary_df.shape[0]+1,
                           cols=location_itinerary_df.shape[1],
                           style='Table Grid')
    # Header of table - center align, add column names and underline
    for j in range(location_itinerary_df.shape[-1]):
        t.cell(0,j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        t.cell(0,j).paragraphs[0].add_run(location_itinerary_df.columns[j]).underline = True
    # Body of table - center align, add date, location, and underline
    for i in range(location_itinerary_df.shape[0]):
        for j in range(location_itinerary_df.shape[-1]):
            t.cell(i+1,j).text = str(location_itinerary_df.values[i,j])
            t.cell(i+1,j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Center align table
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Create month, day, year for current week (M.DD.YYYY)

    # Create title based on particular day of week
    # Monday
    if dt.today().date().weekday() == 0:
        month_day_year = [str((dt.today().date() + timedelta(days=i)).month)+ '.' +\
                          str((dt.today().date() + timedelta(days=i)).day) + '.' +\
                          str((dt.today().date() + timedelta(days=i)).year) for i in range(7, 12, 1)]

    # Tuesday
    if dt.today().date().weekday() == 1:
        month_day_year = [str((dt.today().date() + timedelta(days=i)).month) + '.' + \
                          str((dt.today().date() + timedelta(days=i)).day) + '.' + \
                          str((dt.today().date() + timedelta(days=i)).year) for i in range(6, 11, 1)]

    # Wednesday
    if dt.today().date().weekday() == 2:
        month_day_year = [str((dt.today().date() + timedelta(days=i)).month) + '.' + \
                          str((dt.today().date() + timedelta(days=i)).day) + '.' + \
                          str((dt.today().date() + timedelta(days=i)).year) for i in range(5, 10, 1)]

    # Thursday
    if dt.today().date().weekday() == 3:
        month_day_year = [str((dt.today().date() + timedelta(days=i)).month) + '.' + \
                          str((dt.today().date() + timedelta(days=i)).day) + '.' + \
                          str((dt.today().date() + timedelta(days=i)).year) for i in range(4, 9, 1)]

    # Friday
    if dt.today().date().weekday() == 4:
        month_day_year = [str((dt.today().date() + timedelta(days=i)).month) + '.' + \
                          str((dt.today().date() + timedelta(days=i)).day) + '.' + \
                          str((dt.today().date() + timedelta(days=i)).year) for i in range(3, 8, 1)]

    # Assign title (current date range) to variable
    location_itinerary_word_doc = month_day_year[0] + ' - ' + month_day_year[-1] + ' ' + name_of_user + ' Itinerary.docx'
    # Save created document object
    document.save(location_itinerary_word_doc)
    print('Location itinerary for week of ' + month_day[0] + ' - ' + month_day[-1] + ' complete.')

location_itineraries(name_of_user='Dustin Wicker', work_location=['Denver', 'Denver', 'Denver', 'Denver', 'Denver'])